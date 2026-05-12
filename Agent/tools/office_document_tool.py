"""Чтение и запись Office/PDF: DOCX со стилями Word, простой PDF с типографикой ReportLab, .doc через antiword (если есть).

Стили DOCX — как в Word: Normal, Title, Heading 1…9, Quote, Intense Quote, Subtitle и др. из шаблона документа.
"""
from __future__ import annotations

import json
import re
import subprocess
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from xml.sax.saxutils import escape

from langchain_core.tools import tool

try:
    from ..path_utils import resolve_abs_path
except ImportError:
    from Agent.path_utils import resolve_abs_path

_STYLE_DOC = (
    "Стили абзацев как в Word: Normal, Title, Heading 1 … Heading 9, Subtitle, "
    "Quote, Intense Quote, Caption, List Paragraph. Для PDF-секций role: title, h1, h2, body."
)


def _compact_read(path: str, paragraphs: List[Dict[str, Any]], total: int, extra: str = "") -> str:
    lines = [f"[office_document_read] {path}", f"абзацев в выборке: {len(paragraphs)} из {total}", extra, "---"]
    for p in paragraphs[:40]:
        lines.append(f"[{p.get('i')}] ({p.get('style')}) {str(p.get('text', ''))[:500]}")
    if len(paragraphs) > 40:
        lines.append(f"… ещё {len(paragraphs) - 40} абзацев …")
    return "\n".join(lines)


def _read_doc_antiword(path: Path) -> Tuple[str, str]:
    try:
        r = subprocess.run(
            ["antiword", "-m", "UTF-8.txt", str(path)],
            capture_output=True,
            text=True,
            timeout=120,
        )
        if r.returncode == 0:
            return r.stdout or "", ""
        return "", (r.stderr or f"antiword exit {r.returncode}")[:400]
    except FileNotFoundError:
        return "", "antiword не найден (установите пакет antiword для чтения .doc)"
    except Exception as e:
        return "", str(e)[:400]


def _read_docx_paragraphs(path: Path, max_paragraphs: int) -> Tuple[List[Dict[str, Any]], int, str]:
    try:
        from docx import Document
    except ImportError:
        return [], 0, "python-docx не установлен: pip install python-docx"

    doc = Document(str(path))
    total = len(doc.paragraphs)
    out: List[Dict[str, Any]] = []
    for i, p in enumerate(doc.paragraphs[: max(1, max_paragraphs)]):
        try:
            st = p.style.name if p.style is not None else "Default"
        except Exception:
            st = "Default"
        out.append({"i": i, "style": st, "text": p.text or ""})
    return out, total, ""


def _read_pdf_pages(path: Path, max_pages: int, max_chars: int) -> Tuple[List[Dict[str, Any]], int, str]:
    try:
        import fitz
    except ImportError:
        return [], 0, "pymupdf не установлен: pip install pymupdf"

    doc = fitz.open(str(path))
    n = min(len(doc), max(1, max_pages))
    out: List[Dict[str, Any]] = []
    used = 0
    for pi in range(n):
        text = doc.load_page(pi).get_text("text") or ""
        chunk = text[: max_chars - used] if max_chars > used else ""
        used += len(chunk)
        out.append({"i": pi, "style": f"Page {pi + 1}", "text": chunk})
        if used >= max_chars:
            break
    doc.close()
    return out, len(doc), ""


@tool
def office_document_read(
    file_path: str,
    max_paragraphs: int = 200,
    max_chars: int = 120_000,
) -> Dict[str, Any]:
    """Читает .docx (абзацы + имена стилей Word), .pdf (текст по страницам) или .doc (только текст, нужен antiword).
    Для правок стилей и текста используй docx_document_patch_paragraphs / docx_document_append_paragraphs."""
    path = resolve_abs_path(file_path)
    if not path.is_file():
        return {"error": "not_found", "path": str(path)}

    suf = path.suffix.lower()
    max_paragraphs = max(1, min(int(max_paragraphs), 2000))
    max_chars = max(2000, min(int(max_chars), 500_000))

    paragraphs: List[Dict[str, Any]] = []
    total = 0
    hint = ""

    if suf == ".docx":
        paragraphs, total, err = _read_docx_paragraphs(path, max_paragraphs)
        if err:
            return {"error": "import", "detail": err, "path": str(path)}
        # trim by max_chars
        acc = 0
        trimmed: List[Dict[str, Any]] = []
        for p in paragraphs:
            t = p.get("text") or ""
            if acc + len(t) > max_chars:
                room = max_chars - acc
                if room > 50:
                    trimmed.append({**p, "text": t[:room] + "…"})
                break
            trimmed.append(p)
            acc += len(t)
        paragraphs = trimmed
    elif suf == ".pdf":
        pages = max(1, min(max_paragraphs, 500))
        paragraphs, total, err = _read_pdf_pages(path, pages, max_chars)
        if err:
            return {"error": "import", "detail": err, "path": str(path)}
    elif suf == ".doc":
        raw, err = _read_doc_antiword(path)
        if err and not raw.strip():
            return {
                "path": str(path),
                "error": "doc_read_failed",
                "detail": err,
                "hint": "Конвертируйте .doc → .docx (LibreOffice) или установите antiword.",
            }
        # одна псевдо-страница
        paragraphs = [{"i": 0, "style": "Plain", "text": raw[:max_chars]}]
        total = 1
        hint = "Только текст без стилей (.doc). Для стилей сохраните как .docx."
    else:
        return {
            "error": "unsupported",
            "path": str(path),
            "hint": "Поддерживаются .docx, .pdf, .doc (через antiword).",
        }

    full_text = "\n\n".join((p.get("text") or "") for p in paragraphs)
    compact = _compact_read(str(path), paragraphs, total, hint)
    return {
        "path": str(path),
        "format": suf.lstrip("."),
        "paragraphs": paragraphs,
        "paragraph_count_sampled": len(paragraphs),
        "paragraph_count_total": total,
        "full_text": full_text[: min(len(full_text), max_chars)],
        "hint": hint or _STYLE_DOC,
        "_model_compact": compact,
    }


def _resolve_paragraph_style(doc: Any, name: str) -> Any:
    try:
        return doc.styles[name]
    except Exception:
        try:
            return doc.styles["Normal"]
        except Exception:
            return doc.paragraphs[0].style if doc.paragraphs else None


def _add_block(doc: Any, text: str, style: str) -> None:
    style = (style or "Normal").strip()
    if style.lower() in ("title",):
        doc.add_heading(text, level=0)
        return
    m = re.match(r"heading\s*(\d+)", style, re.I)
    if m:
        lvl = min(9, max(1, int(m.group(1))))
        doc.add_heading(text, level=lvl)
        return
    try:
        doc.add_paragraph(text, style=style)
    except Exception:
        doc.add_paragraph(text, style="Normal")


def _parse_json_list(raw: str, label: str) -> Tuple[Optional[List[Dict[str, Any]]], Optional[str]]:
    try:
        data = json.loads(raw)
    except json.JSONDecodeError as e:
        return None, f"Невалидный JSON ({label}): {e}"
    if not isinstance(data, list):
        return None, f"{label}: ожидается JSON-массив объектов"
    return data, None


@tool
def docx_document_create(file_path: str, paragraphs_json: str) -> Dict[str, Any]:
    """Создаёт новый .docx. paragraphs_json: [{"text":"...","style":"Title|Heading 1|Normal|..."}, ...]."""
    try:
        from docx import Document
    except ImportError:
        return {"error": "import", "detail": "pip install python-docx"}

    items, err = _parse_json_list(paragraphs_json, "paragraphs_json")
    if err:
        return {"error": "bad_json", "detail": err}

    path = resolve_abs_path(file_path)
    if path.suffix.lower() != ".docx":
        path = path.with_suffix(".docx")
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    for it in items:
        if not isinstance(it, dict):
            continue
        text = str(it.get("text", ""))
        style = str(it.get("style", "Normal"))
        if text or style != "Normal":
            _add_block(doc, text, style)

    doc.save(str(path))
    return {"path": str(path), "action": "created", "paragraphs_written": len(items)}


@tool
def docx_document_append_paragraphs(file_path: str, paragraphs_json: str) -> Dict[str, Any]:
    """Добавляет в конец существующего .docx абзацы. paragraphs_json как в docx_document_create."""
    try:
        from docx import Document
    except ImportError:
        return {"error": "import", "detail": "pip install python-docx"}

    path = resolve_abs_path(file_path)
    if not path.is_file():
        return {"error": "not_found", "path": str(path)}
    if path.suffix.lower() != ".docx":
        return {"error": "not_docx", "path": str(path)}

    items, err = _parse_json_list(paragraphs_json, "paragraphs_json")
    if err:
        return {"error": "bad_json", "detail": err}

    doc = Document(str(path))
    n0 = len(doc.paragraphs)
    for it in items:
        if not isinstance(it, dict):
            continue
        text = str(it.get("text", ""))
        style = str(it.get("style", "Normal"))
        _add_block(doc, text, style)
    doc.save(str(path))
    return {
        "path": str(path),
        "action": "appended",
        "paragraphs_before": n0,
        "paragraphs_after": len(doc.paragraphs),
    }


@tool
def docx_document_patch_paragraphs(file_path: str, patches_json: str) -> Dict[str, Any]:
    """Правка абзацев по индексу (0-based). patches_json:
    [{"paragraph_index": 0, "text": "новый текст", "style": "Heading 2"}, ...]
    Поле style можно опустить — останется прежний стиль (кроме смены через text)."""
    try:
        from docx import Document
    except ImportError:
        return {"error": "import", "detail": "pip install python-docx"}

    path = resolve_abs_path(file_path)
    if not path.is_file():
        return {"error": "not_found", "path": str(path)}
    if path.suffix.lower() != ".docx":
        return {"error": "not_docx", "path": str(path)}

    patches, err = _parse_json_list(patches_json, "patches_json")
    if err:
        return {"error": "bad_json", "detail": err}

    doc = Document(str(path))
    total = len(doc.paragraphs)
    applied = 0
    for p in patches:
        if not isinstance(p, dict):
            continue
        try:
            idx = int(p.get("paragraph_index", p.get("i", -1)))
        except (TypeError, ValueError):
            continue
        if idx < 0 or idx >= total:
            return {
                "error": "index_out_of_range",
                "paragraph_index": idx,
                "total_paragraphs": total,
            }
        para = doc.paragraphs[idx]
        if "text" in p:
            para.text = str(p.get("text", ""))
        st = p.get("style")
        if st:
            try:
                para.style = _resolve_paragraph_style(doc, str(st))
            except Exception:
                pass
        applied += 1

    doc.save(str(path))
    return {"path": str(path), "action": "patched", "patches_applied": applied, "paragraphs_total": total}


def _apply_one_docx_op(doc: Any, op: Dict[str, Any]) -> None:
    """Одна операция для docx_document_advanced_ops (python-docx)."""
    name = str(op.get("op", "")).strip().lower()
    if not name:
        raise ValueError("op: пустое имя")

    if name == "append_paragraph":
        text = str(op.get("text", ""))
        style = str(op.get("style", "Normal"))
        _add_block(doc, text, style)
        return

    if name == "set_paragraph_alignment":
        idx = int(op["paragraph_index"])
        align = str(op.get("alignment", "left")).lower()
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        m = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        p = doc.paragraphs[idx]
        p.alignment = m.get(align, WD_ALIGN_PARAGRAPH.LEFT)
        return

    if name == "set_paragraph_spacing":
        idx = int(op["paragraph_index"])
        p = doc.paragraphs[idx]
        pf = p.paragraph_format
        if "space_before_pt" in op:
            from docx.shared import Pt

            pf.space_before = Pt(float(op["space_before_pt"]))
        if "space_after_pt" in op:
            from docx.shared import Pt

            pf.space_after = Pt(float(op["space_after_pt"]))
        rule = str(op.get("line_rule", "")).lower()
        if rule:
            from docx.enum.text import WD_LINE_SPACING
            from docx.shared import Pt

            if rule in ("single", "1"):
                pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            elif rule in ("1.5", "one_point_five", "onepointfive"):
                pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            elif rule in ("double", "2"):
                pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            elif rule == "exact" and "line_spacing_pt" in op:
                pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                pf.line_spacing = Pt(float(op["line_spacing_pt"]))
            elif rule == "multiple" and "line_spacing" in op:
                pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                pf.line_spacing = float(op["line_spacing"])
        return

    if name == "set_paragraph_indent":
        idx = int(op["paragraph_index"])
        p = doc.paragraphs[idx]
        pf = p.paragraph_format
        from docx.shared import Cm

        if "first_line_indent_cm" in op:
            pf.first_line_indent = Cm(float(op["first_line_indent_cm"]))
        if "left_indent_cm" in op:
            pf.left_indent = Cm(float(op["left_indent_cm"]))
        if "right_indent_cm" in op:
            pf.right_indent = Cm(float(op["right_indent_cm"]))
        return

    if name == "set_run_font":
        idx = int(op["paragraph_index"])
        p = doc.paragraphs[idx]
        run_idx = int(op.get("run_index", 0))
        runs = list(p.runs)
        if not runs:
            raise ValueError("в абзаце нет runs")
        if run_idx < 0:
            targets = runs
        elif 0 <= run_idx < len(runs):
            targets = [runs[run_idx]]
        else:
            raise ValueError(f"run_index вне диапазона 0..{len(runs)-1} или -1")
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_UNDERLINE

        for run in targets:
            if "bold" in op:
                run.bold = bool(op["bold"])
            if "italic" in op:
                run.italic = bool(op["italic"])
            if "underline" in op:
                run.underline = WD_UNDERLINE.SINGLE if op["underline"] else WD_UNDERLINE.NONE
            if op.get("font_name"):
                run.font.name = str(op["font_name"])
            if "font_size_pt" in op:
                run.font.size = Pt(float(op["font_size_pt"]))
            hx = str(op.get("color_hex", "")).strip().lstrip("#")
            if len(hx) == 6:
                run.font.color.rgb = RGBColor(
                    int(hx[0:2], 16), int(hx[2:4], 16), int(hx[4:6], 16)
                )
        return

    if name == "set_section_margins":
        si = int(op.get("section_index", 0))
        sec = doc.sections[si]
        from docx.shared import Cm

        for key, attr in (
            ("top_cm", "top_margin"),
            ("bottom_cm", "bottom_margin"),
            ("left_cm", "left_margin"),
            ("right_cm", "right_margin"),
        ):
            if key in op:
                setattr(sec, attr, Cm(float(op[key])))
        return

    if name == "set_section_orientation":
        from docx.enum.section import WD_ORIENT

        si = int(op.get("section_index", 0))
        sec = doc.sections[si]
        ori = str(op.get("orientation", "portrait")).lower()
        if ori == "landscape":
            if sec.orientation != WD_ORIENT.LANDSCAPE:
                w, h = sec.page_width, sec.page_height
                sec.orientation = WD_ORIENT.LANDSCAPE
                sec.page_width, sec.page_height = h, w
        else:
            if sec.orientation != WD_ORIENT.PORTRAIT:
                w, h = sec.page_width, sec.page_height
                sec.orientation = WD_ORIENT.PORTRAIT
                sec.page_width, sec.page_height = h, w
        return

    if name == "set_section_page_size_cm":
        si = int(op.get("section_index", 0))
        sec = doc.sections[si]
        from docx.shared import Cm

        if "width_cm" in op:
            sec.page_width = Cm(float(op["width_cm"]))
        if "height_cm" in op:
            sec.page_height = Cm(float(op["height_cm"]))
        return

    if name == "insert_page_break_after_paragraph":
        idx = int(op["paragraph_index"])
        p = doc.paragraphs[idx]
        from docx.enum.text import WD_BREAK

        p.add_run().add_break(WD_BREAK.PAGE)
        return

    if name == "insert_table_after_paragraph":
        idx = int(op["paragraph_index"])
        rows = max(1, int(op.get("rows", 1)))
        cols = max(1, int(op.get("cols", 1)))
        p = doc.paragraphs[idx]
        table = doc.add_table(rows=rows, cols=cols)
        tbl_el = table._tbl
        parent = tbl_el.getparent()
        if parent is not None:
            parent.remove(tbl_el)
        p._p.addnext(tbl_el)
        cell_texts = op.get("cell_texts")
        if isinstance(cell_texts, list):
            for i, row in enumerate(table.rows):
                if i >= len(cell_texts):
                    break
                row_vals = cell_texts[i]
                if not isinstance(row_vals, list):
                    continue
                for j, cell in enumerate(row.cells):
                    if j >= len(row_vals):
                        break
                    cell.text = str(row_vals[j])
        return

    raise ValueError(f"неизвестная op: {name}")


@tool
def docx_document_advanced_ops(file_path: str, operations_json: str) -> Dict[str, Any]:
    """Массив op в operations_json (до 40): append_paragraph, set_paragraph_*, set_run_font, set_section_*, insert_page_break_after_paragraph, insert_table_after_paragraph. Детали полей — из ответа валидации; TOC/PAGE — через code_interpreter."""
    try:
        from docx import Document
    except ImportError:
        return {"error": "import", "detail": "pip install python-docx"}

    items, err = _parse_json_list(operations_json, "operations_json")
    if err:
        return {"error": "bad_json", "detail": err}
    if len(items) > 40:
        return {"error": "too_many_ops", "max": 40, "got": len(items)}

    path = resolve_abs_path(file_path)
    if not path.is_file():
        return {"error": "not_found", "path": str(path)}
    if path.suffix.lower() != ".docx":
        return {"error": "not_docx", "path": str(path)}

    doc = Document(str(path))
    applied = 0
    errors: List[str] = []
    for i, raw in enumerate(items):
        if not isinstance(raw, dict):
            errors.append(f"#{i}: не объект")
            continue
        try:
            _apply_one_docx_op(doc, raw)
            applied += 1
        except Exception as e:
            errors.append(f"#{i} ({raw.get('op')}): {e}")
    doc.save(str(path))
    return {
        "path": str(path),
        "action": "docx_advanced",
        "applied": applied,
        "errors": errors,
        "ok": not errors,
    }


@tool
def pdf_styled_document_create(file_path: str, sections_json: str, title: str = "") -> Dict[str, Any]:
    """PDF через ReportLab; sections_json: [{role: title|h1|h2|body, text}, ...]; перезапись файла."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError:
        return {"error": "import", "detail": "pip install reportlab"}

    items, err = _parse_json_list(sections_json, "sections_json")
    if err:
        return {"error": "bad_json", "detail": err}

    path = resolve_abs_path(file_path)
    if path.suffix.lower() != ".pdf":
        path = path.with_suffix(".pdf")
    path.parent.mkdir(parents=True, exist_ok=True)

    styles = getSampleStyleSheet()
    story: list = []
    if title.strip():
        story.append(Paragraph(escape(title.strip()), styles["Title"]))
        story.append(Spacer(1, 0.3 * cm))

    for it in items:
        if not isinstance(it, dict):
            continue
        role = str(it.get("role", "body")).lower().strip()
        raw = str(it.get("text", ""))
        safe = escape(raw).replace("\n", "<br/>")
        if role in ("title",):
            story.append(Paragraph(safe, styles["Title"]))
        elif role in ("h1", "heading1", "heading 1"):
            story.append(Paragraph(safe, styles["Heading1"]))
        elif role in ("h2", "heading2", "heading 2"):
            story.append(Paragraph(safe, styles["Heading2"]))
        elif role in ("h3", "heading3"):
            story.append(Paragraph(safe, styles["Heading3"]))
        elif role in ("quote", "quotation"):
            story.append(Paragraph(safe, styles["Quote"]))
        else:
            story.append(Paragraph(safe, styles["Normal"]))
        story.append(Spacer(1, 0.15 * cm))

    doc = SimpleDocTemplate(str(path), pagesize=A4)
    doc.build(story)
    return {"path": str(path), "action": "pdf_created", "sections": len(items)}
