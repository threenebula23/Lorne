"""Инструменты правки .docx с сохранением форматирования (docxedit + python-docx).

См. https://pypi.org/project/docxedit/ — замены по run-ам, таблицы, поиск строки.
"""
from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, Optional

from langchain_core.tools import tool

try:
    from ..path_utils import resolve_abs_path
except ImportError:
    from Agent.path_utils import resolve_abs_path


def _docxedit_available() -> bool:
    try:
        import docxedit  # noqa: F401
        from docx import Document  # noqa: F401
        return True
    except ImportError:
        return False


@tool
def docxedit_replace_keep_format(
    file_path: str,
    old_string: str,
    new_string: str,
    include_tables: bool = True,
) -> Dict[str, Any]:
    """Заменить все вхождения old_string на new_string в .docx, сохраняя стили (docxedit.replace_string).

    Args:
        file_path: Путь к .docx
        old_string: Текст для поиска
        new_string: Замена
        include_tables: учитывать текст в таблицах (как в docxedit)
    """
    if not _docxedit_available():
        return {"error": "import", "detail": "pip install docxedit python-docx"}
    import docxedit
    from docx import Document

    path = resolve_abs_path(file_path)
    if not path.is_file():
        return {"error": "not_found", "path": str(path)}
    if path.suffix.lower() != ".docx":
        return {"error": "not_docx", "path": str(path)}
    if not old_string:
        return {"error": "empty_old_string"}

    doc = Document(str(path))
    try:
        docxedit.replace_string(
            doc,
            old_string=old_string,
            new_string=new_string,
            include_tables=include_tables,
        )
    except TypeError:
        docxedit.replace_string(doc, old_string=old_string, new_string=new_string)
    doc.save(str(path))
    return {"path": str(path), "action": "docxedit_replace", "ok": True}


@tool
def docxedit_replace_up_to_paragraph(
    file_path: str,
    old_string: str,
    new_string: str,
    paragraph_number: int,
    include_tables: bool = True,
) -> Dict[str, Any]:
    """Замена только до указанного номера абзаца (1-based как в docxedit.replace_string_up_to_paragraph)."""
    if not _docxedit_available():
        return {"error": "import", "detail": "pip install docxedit python-docx"}
    import docxedit
    from docx import Document

    path = resolve_abs_path(file_path)
    if not path.is_file():
        return {"error": "not_found", "path": str(path)}
    if path.suffix.lower() != ".docx":
        return {"error": "not_docx"}
    doc = Document(str(path))
    try:
        docxedit.replace_string_up_to_paragraph(
            doc,
            old_string=old_string,
            new_string=new_string,
            paragraph_number=int(paragraph_number),
            include_tables=include_tables,
        )
    except TypeError:
        docxedit.replace_string_up_to_paragraph(
            doc, old_string=old_string, new_string=new_string,
            paragraph_number=int(paragraph_number),
        )
    doc.save(str(path))
    return {"path": str(path), "action": "docxedit_replace_limited", "ok": True}


@tool
def docxedit_find_line(file_path: str, search_text: str) -> Dict[str, Any]:
    """Показать строку/контекст, где найден search_text (docxedit.show_line)."""
    if not _docxedit_available():
        return {"error": "import", "detail": "pip install docxedit python-docx"}
    import docxedit
    from docx import Document

    path = resolve_abs_path(file_path)
    if not path.is_file():
        return {"error": "not_found", "path": str(path)}
    doc = Document(str(path))
    try:
        line = docxedit.show_line(doc, current_text=search_text)
    except Exception as e:
        return {"error": str(e), "path": str(path)}
    return {"path": str(path), "line": str(line) if line is not None else ""}


@tool
def docxedit_table_cell_append(
    file_path: str,
    table_index: int,
    row_num: int,
    column_num: int,
    new_string: str,
) -> Dict[str, Any]:
    """Добавить текст в ячейку таблицы (docxedit.add_text_in_table). Индексы таблицы 0-based; row/column как в docxedit (1-based в API)."""
    if not _docxedit_available():
        return {"error": "import", "detail": "pip install docxedit python-docx"}
    import docxedit
    from docx import Document

    path = resolve_abs_path(file_path)
    if not path.is_file():
        return {"error": "not_found"}
    doc = Document(str(path))
    tables = doc.tables
    if table_index < 0 or table_index >= len(tables):
        return {"error": "bad_table_index", "tables": len(tables)}
    table = tables[table_index]
    try:
        docxedit.add_text_in_table(
            table,
            row_num=int(row_num),
            column_num=int(column_num),
            new_string=new_string,
        )
    except Exception as e:
        return {"error": str(e)}
    doc.save(str(path))
    return {"path": str(path), "action": "docxedit_table_cell", "ok": True}


@tool
def docxedit_table_font_size(file_path: str, table_index: int, font_size: int) -> Dict[str, Any]:
    """Задать размер шрифта для всей таблицы (pt), docxedit.change_table_font_size."""
    if not _docxedit_available():
        return {"error": "import", "detail": "pip install docxedit python-docx"}
    import docxedit
    from docx import Document

    path = resolve_abs_path(file_path)
    if not path.is_file():
        return {"error": "not_found"}
    doc = Document(str(path))
    if table_index < 0 or table_index >= len(doc.tables):
        return {"error": "bad_table_index", "tables": len(doc.tables)}
    table = doc.tables[table_index]
    try:
        docxedit.change_table_font_size(table, font_size=int(font_size))
    except Exception as e:
        return {"error": str(e)}
    doc.save(str(path))
    return {"path": str(path), "action": "docxedit_table_font", "ok": True}
